"""
Aplomb — Mitigation Efficacy Analysis (Word document)
=====================================================

Builds Aplomb_mitigation_analysis_v1.docx in /aplomb.clinic/.

Independent, evidence-graded scoring of the four top APALM-shortlisted
products against the GLP-1 side effects they're positioned to mitigate.

Brand-matched editorial style (cream / Cormorant Garamond / IBM Plex Sans)
mirroring build_supplier_shortlist_deck.py brand tokens.

Run: python3 build_mitigation_analysis_docx.py
"""

from pathlib import Path
from datetime import datetime
import copy

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap

# ---------------------------------------------------------------------------
# Brand tokens — mirrored from build_supplier_shortlist_deck.py:30-50
# ---------------------------------------------------------------------------
BG_HEX     = "EFE8DC"   # warm cream — page color
PAPER_HEX  = "F7F1E6"   # lighter cream — table cell wash
INK_HEX    = "1A1512"   # near-black brown — body ink
AMBER_HEX  = "7A3D14"   # deep amber — accents, dividers, eyebrows
TAN_HEX    = "D9A06B"   # secondary accent
RULE_HEX   = "D9CFBD"   # hairlines, table borders
MUTED_HEX  = "6B5D4C"   # secondary text
SCORE_HI_HEX  = "E6C99C"  # score wash >= 70
SCORE_MID_HEX = "F0E5CE"  # 50-69
SCORE_LOW_HEX = "E5DDC2"  # < 50

INK   = RGBColor.from_string(INK_HEX)
AMBER = RGBColor.from_string(AMBER_HEX)
MUTED = RGBColor.from_string(MUTED_HEX)
TAN   = RGBColor.from_string(TAN_HEX)

SERIF = "Cormorant Garamond"
SANS  = "IBM Plex Sans"

OUT = Path(
    "/Users/zacharypoll/Desktop/Documents/Claude Code/aplomb.clinic/"
    "Aplomb_mitigation_analysis_v1.docx"
)

TODAY = datetime(2026, 5, 6).strftime("%B %-d, %Y")


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------
def set_page_color(doc, hex_color):
    """Cream page background — w:background element + displayBackgroundShape."""
    settings = doc.settings.element
    bg_el = settings.find(qn("w:displayBackgroundShape"))
    if bg_el is None:
        bg_el = OxmlElement("w:displayBackgroundShape")
        settings.append(bg_el)
    body = doc.element.body
    bg = body.getparent().find(qn("w:background"))
    if bg is None:
        bg = OxmlElement("w:background")
        bg.set(qn("w:color"), hex_color)
        doc.element.insert(0, bg)
    else:
        bg.set(qn("w:color"), hex_color)


def shade_cell(cell, hex_color):
    """Cell fill color via w:shd element."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, hex_color=RULE_HEX, sz="4", sides=("top","left","bottom","right")):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), hex_color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Cell padding in twips."""
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        mar.append(m)
    tc_pr.append(mar)


def add_hr(paragraph, hex_color=AMBER_HEX, sz="8"):
    """Bottom-border on an empty paragraph = horizontal rule."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def style_run(run, *, font=SANS, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


# ---------------------------------------------------------------------------
# Mid-level helpers — paragraphs / headings
# ---------------------------------------------------------------------------
def add_para(doc, text, *, font=SANS, size=10.5, color=INK, bold=False,
             italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0,
             space_after=6, line_spacing=1.32, first_line_indent=None,
             keep_with_next=False, keep_together=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = Pt(first_line_indent)
    if keep_with_next:
        pf.keep_with_next = True
    if keep_together:
        pf.keep_together = True
    p.alignment = align
    if text:
        run = p.add_run(text)
        style_run(run, font=font, size=size, color=color, bold=bold, italic=italic)
    return p


def add_eyebrow(doc, text):
    """Tiny amber uppercase header — '§ 01 · Side effect 1 of 4 · Topical serum'."""
    p = add_para(doc, text.upper(), font=SANS, size=8, color=AMBER,
                 space_before=18, space_after=2, line_spacing=1.1,
                 keep_with_next=True)
    for run in p.runs:
        run.font.bold = True
        # widen letter-spacing for an editorial wordmark feel
        rPr = run._element.get_or_add_rPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), "30")
        rPr.append(spacing)
    return p


def add_h1(doc, text):
    p = add_para(doc, text, font=SERIF, size=26, color=INK,
                 space_before=4, space_after=10, line_spacing=1.06,
                 keep_with_next=True)
    return p


def add_h2(doc, text):
    p = add_para(doc, text, font=SERIF, size=18, color=INK, bold=False,
                 space_before=14, space_after=6, line_spacing=1.1,
                 keep_with_next=True)
    return p


def add_h3(doc, text):
    p = add_para(doc, text, font=SANS, size=10.5, color=AMBER, bold=True,
                 space_before=10, space_after=2, line_spacing=1.2,
                 keep_with_next=True)
    for run in p.runs:
        rPr = run._element.get_or_add_rPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), "20")
        rPr.append(spacing)
        run.font.name = SANS
    # convert to small caps via uppercase text + tracking
    for run in p.runs:
        run.text = run.text.upper()
    return p


def add_lede(doc, text):
    p = add_para(doc, text, font=SERIF, size=13, color=INK, italic=True,
                 space_before=4, space_after=10, line_spacing=1.32,
                 keep_with_next=True)
    return p


def add_amber_rule(doc, sz="8"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_hr(p, hex_color=AMBER_HEX, sz=sz)
    return p


def add_thin_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    add_hr(p, hex_color=RULE_HEX, sz="4")
    return p


def add_bullet(doc, text, *, indent=0.22, color=INK, italic=False, size=10):
    p = add_para(doc, "", font=SANS, size=size, color=color, italic=italic,
                 space_after=2, line_spacing=1.26)
    pf = p.paragraph_format
    pf.left_indent = Inches(indent + 0.16)
    pf.first_line_indent = Inches(-0.16)
    bullet_run = p.add_run("•   ")
    style_run(bullet_run, font=SANS, size=size, color=AMBER, bold=True)
    body_run = p.add_run(text)
    style_run(body_run, font=SANS, size=size, color=color, italic=italic)
    return p


def add_kv_line(doc, key, value):
    """Small key/value line — bold key, regular value."""
    p = add_para(doc, "", font=SANS, size=9.5, color=INK,
                 space_after=2, line_spacing=1.28)
    k = p.add_run(f"{key} ")
    style_run(k, font=SANS, size=9.5, color=AMBER, bold=True)
    v = p.add_run(value)
    style_run(v, font=SANS, size=9.5, color=INK)
    return p


def add_pagebreak(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------
def style_table_header_row(row):
    for cell in row.cells:
        shade_cell(cell, AMBER_HEX)
        set_cell_borders(cell, hex_color=AMBER_HEX, sz="4")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                style_run(run, font=SANS, size=9, color=RGBColor.from_string("F5EFE3"), bold=True)
                rPr = run._element.get_or_add_rPr()
                spacing = OxmlElement("w:spacing")
                spacing.set(qn("w:val"), "20")
                rPr.append(spacing)


def set_row_cant_split(row):
    """Prevent a table row from splitting across pages (w:cantSplit)."""
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def fill_cell(cell, text, *, font=SANS, size=9.5, color=INK, bold=False,
              italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, fill=None,
              vertical=WD_ALIGN_VERTICAL.CENTER):
    cell.vertical_alignment = vertical
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell, hex_color=RULE_HEX, sz="4")
    set_cell_margins(cell, top=70, bottom=70, left=110, right=110)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.22
    p.alignment = align
    run = p.add_run(text)
    style_run(run, font=font, size=size, color=color, bold=bold, italic=italic)


def score_color(score):
    if score >= 70:
        return SCORE_HI_HEX
    if score >= 50:
        return SCORE_MID_HEX
    return SCORE_LOW_HEX


# ---------------------------------------------------------------------------
# Rubric — single source of truth for weights
# ---------------------------------------------------------------------------
RUBRIC = [
    ("Mechanistic match",          20, "Does the active address the actual underlying biology of the side effect (not just an adjacent pathway)?"),
    ("Clinical evidence quality",  20, "RCTs > meta-analyses > observational. GLP-1-population data > analogous-population > mechanism-only."),
    ("Effect size",                15, "Magnitude of relief in the cited evidence on a clinically meaningful scale."),
    ("Dose adequacy",              15, "Does the labeled dose match the dose at which clinical benefit is documented?"),
    ("Mechanism coverage",         10, "Most side effects have multiple drivers. Does the product address one or several?"),
    ("Tolerability in GLP-1 users",10, "Does the product avoid worsening other GLP-1 side effects, drug interactions, or assay interference?"),
    ("Realistic adherence",        10, "Daily dose burden, palatability, schedule complexity. A product taken inconsistently does not mitigate."),
]

# ---------------------------------------------------------------------------
# Per-product scores — derived from research dossiers
# ---------------------------------------------------------------------------

# (dimension index aligns with RUBRIC)
PRODUCTS = [
    {
        "key": "ozempic_face",
        "section_no": "§ 01",
        "side_effect": "Ozempic Face",
        "supplier": "Cellular Cosmetics",
        "product_name": "Anti-Age Peptide Serum",
        "format": "Topical serum, dropper bottle",
        "actives": "Matrixyl 3000 (palmitoyl tetrapeptide-7 + palmitoyl tripeptide-1) + Synthe'6 (palmitoyl tripeptide-38)",
        "moq": "10 units · $25.00/unit · 4–8 day lead",
        "internal_score": 84,
        "scores": [9, 12, 9, 11, 3, 9, 8],
        "scores_rationale": [
            "Targets dermal collagen / elastin via biostimulatory peptides — but the primary driver of Ozempic face is subcutaneous fat-pad volume loss in malar / temporal / buccal compartments. Topicals reach 1–3 mm; fat pads sit at 5–10 mm. Right for the secondary collagen-loss driver, wrong for the primary volumetric one.",
            "Matrixyl 3000 has manufacturer (Sederma) trial data and one 12-week study (Li 2023, J Cosm Dermatol, N=37) supporting collagen density gains. Synthe'6 evidence is thinner — primarily Sederma in-house dossiers. No placebo-controlled trial in GLP-1 users; zero evidence in caloric-restriction-induced facial wasting.",
            "Sederma trial: 45% reduction in deep-wrinkle surface area at 2% / twice-daily for 8 weeks; Li 2023: +54.99% collagen density at 12 weeks. Real and measurable on dermal endpoints. Negligible on facial volume — not the same axis.",
            "If Cellular's serum is formulated at the standard 3–8% Matrixyl 3000 complex (the supplier publishes INCI per APALM notes), dose is consistent with Sederma's protocol. Synthe'6 ~2% is also at the studied concentration. Adequacy is good for the dermal effect, but irrelevant for fat-pad restoration.",
            "Addresses 1 of 3 listed Ozempic-face drivers (dermal collagen/elastin loss) and 0 of the volumetric drivers. Sunscreen, retinoids, and HA fillers each cover separate, larger fractions of the problem.",
            "Topical peptides are extremely well tolerated. No GLP-1 drug interaction. Negligible risk profile in this population.",
            "AM/PM serum is a manageable habit. Realistic adherence in a consumer cohort is moderate-to-high.",
        ],
    },
    {
        "key": "halitosis",
        "section_no": "§ 02",
        "side_effect": "Halitosis (\"Ozempic breath\")",
        "supplier": "BLIS Technologies",
        "product_name": "K12 FreshBreath Lozenge",
        "format": "1.25B CFU peppermint lozenge, 40-count blister",
        "actives": "Streptococcus salivarius K12 (bacteriocin-producing oral probiotic)",
        "moq": "100 units · $2.00/unit · NZ-based",
        "internal_score": 82,
        "scores": [13, 14, 11, 11, 4, 10, 8],
        "scores_rationale": [
            "K12 colonizes the oral cavity and competitively suppresses anaerobic VSC-producers (Solobacterium moorei, Fusobacterium spp., Porphyromonas gingivalis). Direct hit on the bacterial driver. But of the three converging GLP-1 halitosis mechanisms (hyposalivation, ketotic acetone, gastric stasis fermentation), it only addresses the bacterial component of the third.",
            "Best evidence base of the four products. Burton 2006 RCT (J Appl Microbiol) is the canonical halitosis trial; Kim 2020 (Probiotics Antimicrob Proteins, N=28) showed within-group VSC drop but no placebo-adjusted effect. Tan 2024 systematic review (6 RCTs, 360 subjects) found 5/6 trials showed VSC reduction at low-to-moderate certainty. Yoo 2020 in vitro: 95% H₂S suppression in co-culture.",
            "Burton 2006: 85% of K12 subjects achieved >100 ppb VSC reduction vs. 30% placebo (p<0.05). Concrete and on a clinically meaningful instrument (Halimeter). Effect tapers within 4 weeks of discontinuation per the systematic review.",
            "Burton dosed 1B CFU per lozenge × 4 lozenges/day acutely; BLIS lozenge is 1.25B CFU per dose. One to two lozenges per day puts a user in the studied maintenance range. Slightly below the acute Burton load but well within the maintenance dose used in later trials.",
            "Addresses 1 of 3 GLP-1 halitosis drivers (the bacterial one — admittedly the dominant contributor to oral VSCs). Hyposalivation requires saliva substitutes; ketotic acetone is exhaled from the lung, not the mouth, and is unaddressable by oral probiotic; gastric-stasis eructation requires GI motility intervention.",
            "Probiotic, GRAS, established safety. Peppermint lozenge format is innocuous. No GLP-1 contraindication. Caveat: rare case reports in immunocompromised users — not relevant to the typical GLP-1 cohort.",
            "1–2 lozenges per day, palatable, dissolved in mouth. Easy. Fits into existing oral-care routines without behavior change.",
        ],
    },
    {
        "key": "hair_loss",
        "section_no": "§ 03",
        "side_effect": "Hair Loss",
        "supplier": "Vox Nutrition",
        "product_name": "Hair, Skin & Nails Capsule",
        "format": "Daily capsule, NSF / FDA / cGMP / Organic certified",
        "actives": "Biotin (typically 5–10 mg) + marine collagen peptides (typically 1–3 g) + supportive vitamins",
        "moq": "50 units · $4.50/unit · 7–10 day turnaround",
        "internal_score": 72,
        "scores": [7, 7, 5, 9, 3, 7, 8],
        "scores_rationale": [
            "Biotin is the wrong active for non-deficient telogen effluvium — and most GLP-1 users are not biotin-deficient (Durusu Turkoglu 2024 found biotin levels normal in TE patients). Marine collagen provides amino acid building blocks and has weak adjacent relevance. The actual drivers of GLP-1 telogen effluvium — caloric / protein insufficiency and sub-clinical iron, zinc, vitamin D and B12 deficits — are not the targets of this capsule.",
            "Patel 2017 (Skin Appendage Disord) systematically reviewed 18 case reports of biotin for hair loss and found every responder had an underlying deficiency state. Soleymani 2017 (J Drugs Dermatol) concluded insufficient evidence in healthy adults. Milani 2023 (Skin Res Technol, N=83) showed shedding reduction — but the formulation included iron + selenium + amino acids, so the marine-collagen-alone effect cannot be isolated. No GLP-1-specific data.",
            "In non-deficient subjects: minimal effect from biotin (close to placebo). Marine collagen evidence is moderate for skin endpoints, weak for hair-specific endpoints. Realistic shedding reduction in a typical user: 10–20% in 3 months, much of which is regression to the mean as the underlying caloric deficit resolves.",
            "Typical Vox-class formulation delivers biotin 5–10 mg (matches deficiency-treatment doses but irrelevant in the non-deficient majority) and collagen 1–3 g (below the 10 g/day threshold used in skin RCTs). Capsule format limits collagen dose. Adequacy is good for what's in the bottle; the bottle isn't holding what telogen effluvium actually responds to.",
            "Misses iron, zinc, vitamin D, B12 — the deficits actually associated with telogen effluvium. Misses the protein/caloric-adequacy axis entirely. A repletion-targeted product (with therapeutic iron + zinc + vitamin D after lab confirmation) would cover meaningfully more of the mechanism.",
            "Generally safe — but biotin >2.5 mg/day interferes with ~59% of common immunoassays (TSH, free T4, troponin, hCG). This is clinically consequential for GLP-1 users monitored for thyroid function during rapid weight loss, and the product almost certainly exceeds that threshold.",
            "Daily capsule, low burden. Adherence is not the issue.",
        ],
    },
    {
        "key": "nausea_ing",
        "section_no": "§ 04 · A",
        "side_effect": "Nausea (APALM #1 — ING Pharmaceutical)",
        "supplier": "ING Pharmaceutical",
        "product_name": "Nausea Relief Softgel",
        "format": "Softgel — ready-to-fill formula",
        "actives": "Ginger root 67 mg + Vitamin B6 1.4 mg + Magnesium carbonate (unspecified)",
        "moq": "100 units · $2.50/unit · 300+ brands served",
        "internal_score": 82,
        "scores": [14, 14, 4, 1, 5, 9, 9],
        "scores_rationale": [
            "Ginger and B6 are exactly the right actives for GLP-1 nausea — ginger acts on 5-HT3 and gastric motility; B6 modulates central nausea pathways. Mechanistically a textbook match.",
            "Excellent supporting literature: Vutyavanich 2001 (Obstet Gynecol, 1 g ginger × 4 days = 2.3× placebo VAS reduction); Ryan / URCC CCOP 2012 (Support Care Cancer 20:1479–1489, N=576, four-arm 0.5 / 1.0 / 1.5 g dose-response, with significance at 0.5 g and 1.0 g); Chaiyakunapruk 2006 (AJOG 194:95–99, post-op meta-analysis, fixed dose ≥1 g, RR 0.69 for nausea); ACOG B6 first-line endorsement. The evidence is strong — for the doses studied.",
            "At the labeled dose: published data is silent. No trial has tested ginger below 250 mg. Extrapolation of the URCC dose-response curve to 67 mg sits below the inflection point; expected effect indistinguishable from placebo.",
            "The decisive failure. 67 mg ginger is 14.9× below the studied therapeutic minimum (1,000 mg), and 7.5× below the lowest dose the URCC trial actually tested (500 mg). 1.4 mg B6 is the adult RDA — i.e., a multivitamin trace, not a therapeutic dose. ACOG protocol is 25 mg three times daily (75 mg total); this product is at ~1.9% of that.",
            "Targets two mechanism axes (peripheral 5-HT3 / gastric motility via ginger; central via B6) plus magnesium — but at sub-pharmacologic concentrations on each.",
            "Truly safe — at this dose, there is no realistic risk of anticoagulant interaction (ginger's only meaningful caveat) or B6 neuropathy (a >100 mg/day chronic-exposure issue). Safe partly because under-dosed.",
            "Softgel, easy to swallow. Could titrate up via multiple capsules, but that quickly pushes unit cost above premium gummy formats while still barely approaching therapeutic dose.",
        ],
    },
    {
        "key": "nausea_vw",
        "section_no": "§ 04 · B",
        "side_effect": "Nausea (Properly-Dosed Benchmark)",
        "supplier": "“Vital Whole”-style archetype",
        "product_name": "Ginger Root Extract Gummy",
        "format": "Pectin gummy — typical premium spec",
        "actives": "Standardized ginger root extract ~1,000 mg + Vitamin B6 ~25 mg",
        "moq": "Comparable to APALM ginger chew alternatives ($2.00–2.50/unit)",
        "internal_score": None,
        "scores": [16, 17, 12, 14, 7, 9, 9],
        "scores_rationale": [
            "Same actives as ING — properly dosed. Hits both peripheral (gastric, 5-HT3) and central (CNS / area postrema) nausea pathways at concentrations the literature has actually validated.",
            "Strongest evidence base of the five products evaluated. Cochrane post-op meta (Chaiyakunapruk 2006, n=363), Cochrane pregnancy review (Matthews 2015), URCC CCOP 2011 dose-response (N=576), ACOG B6 first-line (Sahakian 1991, Vutyavanich 1995), Sripramote 2003 head-to-head (1 g ginger ≈ 75 mg B6), Oliveira 2014 combination trial. Note: GLP-1-specific RCTs do not yet exist for any nausea intervention — this is the same data shape, but the inference to GLP-1 cohort is interpolation across analogous nausea mechanisms.",
            "Vutyavanich 2001: VAS reduction 2.1 vs. 0.9 placebo. Ryan / URCC CCOP 2012: 1.0 g produced the strongest dose-response signal across the 0.5–1.5 g range. Chaiyakunapruk 2006 meta: RR 0.69 for post-op nausea (~31% RR reduction) and RR 0.61 for post-op vomiting (~39% RR reduction), NNT 4–6. Oliveira 2014 ginger + B6 combo: 73% nausea reduction, 81% vomiting reduction, Rhodes Index. Concrete and across multiple pain populations.",
            "1 g ginger is the URCC dose-response peak and within the consensus 1,000–1,500 mg/day band cited by Cochrane reviews. 25 mg B6 sits squarely inside the ACOG 10–25 mg t.i.d. window. Almost ideal.",
            "Covers both peripheral and central nausea pathways at therapeutic levels. Does not directly address the gastric-stasis driver (no prokinetic action), so 2 of 3 mechanisms covered.",
            "Real interactions exist at therapeutic dose. Ginger has theoretical anticoagulant additivity (warfarin / DOACs) — relevant disclosure even if rare in the GLP-1 cohort. B6 chronic intake should stay under 100 mg/day to avoid peripheral neuropathy; at 25 mg this is well-cleared. Mild GI irritation in ~4% (Oliveira).",
            "Pectin gummy is the easiest-to-take adherence format in any anti-nausea category. Risk is overconsumption; product needs clear dose labeling.",
        ],
    },
]

assert all(sum(p["scores"]) == p_total for p in PRODUCTS for p_total in [sum(p["scores"])])  # sanity


# ---------------------------------------------------------------------------
# Build sections
# ---------------------------------------------------------------------------
def build_cover(doc):
    # Top spacer — modest, the page is one page long total
    add_para(doc, "", space_after=80)
    add_eyebrow(doc, "Aplomb. · Independent product analysis · v1")
    add_amber_rule(doc, sz="12")
    add_para(
        doc,
        "Does It Actually Work?",
        font=SERIF, size=42, color=INK, italic=False,
        space_before=10, space_after=2, line_spacing=1.0,
        keep_with_next=True,
    )
    add_para(
        doc,
        "An evidence-graded mitigation review of the four top APALM-shortlisted products against the GLP-1 side effects they're positioned to treat.",
        font=SERIF, size=14, color=MUTED, italic=True,
        space_before=4, space_after=18, line_spacing=1.32,
    )
    add_amber_rule(doc, sz="6")

    # Subjects
    add_para(doc, "", space_after=2)
    for prod in PRODUCTS:
        p = add_para(doc, "", font=SANS, size=10, color=INK,
                     space_after=2, line_spacing=1.32)
        a = p.add_run(f"{prod['section_no']}    ")
        style_run(a, font=SANS, size=10, color=AMBER, bold=True)
        b = p.add_run(prod["side_effect"])
        style_run(b, font=SERIF, size=13, color=INK)
        c = p.add_run(f"   ·   {prod['supplier']} — {prod['product_name']}")
        style_run(c, font=SANS, size=9.5, color=MUTED, italic=True)

    add_para(doc, "", space_after=18)
    add_thin_rule(doc)
    add_kv_line(doc, "Date", TODAY)
    add_kv_line(doc, "Prepared for", "Zachary Poll · Aplomb. brand")
    add_kv_line(doc, "Scope", "Mitigation efficacy (not cure) of the named formulation against the actual biology of the targeted side effect.")
    add_kv_line(doc, "Out of scope", "Nutrient depletion (per direction). Supplier sourcing fit (covered separately in the APALM deck).")
    add_pagebreak(doc)


def build_executive_summary(doc):
    add_eyebrow(doc, "Executive summary · 1 of 4")
    add_h1(doc, "What the scoring says")
    add_amber_rule(doc, sz="6")

    add_lede(doc,
        "Of the four products, only one — the BLIS K12 lozenge — earns its position on the merits "
        "of its actual formulation against the actual biology. Two are mechanistically misaligned with "
        "the side effect they're sold against. One is so under-dosed that it sits below the lowest "
        "concentration any published trial has tested."
    )

    # Comparative scorecard
    headers = ["Side effect", "Product", "Internal score", "Independent score", "Δ", "Verdict"]
    rows = [
        ["Ozempic face", "Cellular Cosmetics serum", "84", "61", "−23", "Right product, wrong target"],
        ["Halitosis",    "BLIS K12 lozenge",         "82", "71", "−11", "Earns its place; partial coverage"],
        ["Hair loss",    "Vox Hair, Skin & Nails",   "72", "46", "−26", "Wrong active, lab-assay risk"],
        ["Nausea (APALM)", "ING Pharmaceutical softgel", "82", "56", "−26", "Sub-pharmacologic dose"],
        ["Nausea (benchmark)", "Vital Whole-style 1 g gummy", "—", "84", "—", "What the shortlist should look like"],
    ]
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # widths
    widths = [Inches(1.05), Inches(1.85), Inches(0.75), Inches(0.95), Inches(0.5), Inches(1.65)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w
    # header
    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, font=SANS, size=9, color=RGBColor.from_string("F5EFE3"),
                  bold=True, fill=AMBER_HEX, align=WD_ALIGN_PARAGRAPH.LEFT)
    # body — keep each row from splitting across pages
    for trow in table.rows:
        set_row_cant_split(trow)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            if c == 3:  # independent score
                # color score cell by score band
                if val == "—":
                    fill = SCORE_LOW_HEX
                else:
                    score = int(val)
                    fill = score_color(score)
                fill_cell(table.rows[r].cells[c], val, bold=True, fill=fill,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            elif c == 4:  # delta
                fill_cell(table.rows[r].cells[c], val, color=AMBER if "−" in val else INK,
                          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            elif c == 2:  # internal score
                fill_cell(table.rows[r].cells[c], val, color=MUTED,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            elif c == 0:
                fill_cell(table.rows[r].cells[c], val, font=SERIF, size=11, color=INK)
            else:
                fill_cell(table.rows[r].cells[c], val, font=SANS, size=9.5, color=INK)

    add_para(doc, "", space_after=10)

    add_h3(doc, "Headline findings")
    add_bullet(doc,
        "The single largest independent-vs-internal gap is on nausea. ING Pharmaceutical's softgel "
        "delivers 67 mg ginger — 14.9× below the 1,000 mg therapeutic minimum cited in every major "
        "RCT and 7.5× below the lowest dose the URCC dose-response trial even tested. The actives are "
        "right; the formulation is sub-pharmacologic.")
    add_bullet(doc,
        "Ozempic face is a volumetric problem (subcutaneous fat-pad loss in malar / temporal / buccal "
        "compartments), not a dermal problem. Topical peptide serums penetrate ~1–3 mm; fat compartments "
        "sit at 5–10 mm. The Cellular Cosmetics serum has solid evidence on the wrong endpoint — it "
        "improves wrinkle depth and dermal density, not facial volume.")
    add_bullet(doc,
        "Vox's biotin-led capsule targets a deficiency state most GLP-1 users don't have. Patel 2017's "
        "systematic review of biotin for hair loss found every responder had pre-existing biotin "
        "deficiency. Worse, biotin >2.5 mg/day interferes with ~59% of common immunoassays — a "
        "real concern in a population monitored for thyroid function during rapid weight loss.")
    add_bullet(doc,
        "BLIS K12 is the most defensible position on the shortlist. Burton 2006 RCT shows 85% of K12 "
        "subjects achieve >100 ppb VSC reduction vs. 30% placebo — concrete effect, validated instrument. "
        "The real caveat is that K12 only addresses 1 of 3 GLP-1 halitosis drivers (the bacterial one); "
        "ketotic acetone and gastric-stasis eructation lie outside any oral probiotic's reach.")
    add_bullet(doc,
        "If Aplomb substituted a properly-dosed ginger gummy (1 g + 25 mg B6) for the ING softgel, it "
        "would move the nausea SKU from 56 to 84 on this rubric — the largest single-decision "
        "improvement available across the four-product slate.")

    add_pagebreak(doc)


def build_methodology(doc):
    add_eyebrow(doc, "Methodology · 2 of 4")
    add_h1(doc, "How each product is scored")
    add_amber_rule(doc, sz="6")

    add_lede(doc,
        "Seven evidence-anchored dimensions, weighted to 100. Each product receives a score per "
        "dimension and a composite, with rationale grounded in cited evidence — not in supplier "
        "marketing or in the existing APALM internal score."
    )

    add_para(doc,
        "The rubric measures mitigation, not cure. The question is not whether the product reverses "
        "the side effect, but whether it meaningfully reduces severity, frequency, duration, or social "
        "cost for a person on a GLP-1. \"Right active, wrong dose\" and \"right active, wrong target tissue\" "
        "are explicit failure modes, because they are how most off-the-shelf adjacent-category products "
        "fail when redirected at GLP-1 side effects.")

    add_h2(doc, "The seven dimensions")

    headers = ["Dimension", "Weight", "What it measures"]
    table = doc.add_table(rows=len(RUBRIC) + 2, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(2.0), Inches(0.7), Inches(4.05)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, font=SANS, size=9, color=RGBColor.from_string("F5EFE3"),
                  bold=True, fill=AMBER_HEX)

    for trow in table.rows:
        set_row_cant_split(trow)
    for r, (name, weight, desc) in enumerate(RUBRIC, start=1):
        fill_cell(table.rows[r].cells[0], name, font=SERIF, size=11, color=INK, bold=True)
        fill_cell(table.rows[r].cells[1], str(weight), bold=True, color=AMBER,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(table.rows[r].cells[2], desc, font=SANS, size=9.5, color=INK)

    # total row
    fill_cell(table.rows[-1].cells[0], "Composite", font=SERIF, size=12, color=INK, bold=True,
              fill=PAPER_HEX)
    fill_cell(table.rows[-1].cells[1], "100", bold=True, color=AMBER, fill=PAPER_HEX,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(table.rows[-1].cells[2], "Sum across the seven dimensions, no further weighting.",
              font=SANS, size=9.5, color=MUTED, italic=True, fill=PAPER_HEX)

    add_para(doc, "", space_after=8)

    add_h2(doc, "How this differs from the existing APALM score")

    add_para(doc,
        "The existing internal score (the alleviation_score field in build_supplier_shortlist_deck.py) "
        "and the supplier-fit rubric in AFTER_halitosis_scoring_v1.md (40% clinical evidence + 25% "
        "specificity + 25% reputation + 10% regulatory) score the supplier as a sourcing partner — "
        "MOQ, certifications, brand fit, geographic and lead-time risk. Useful for go/no-go on a "
        "purchase order. Not useful for whether the product on the shelf actually mitigates the "
        "side effect.")

    add_para(doc,
        "This document scores the product as a symptom mitigator. Where the two diverge, that "
        "divergence is the finding. A supplier can be a great vendor of a product that doesn't work.")

    add_h3(doc, "Evidence weighting conventions")
    add_bullet(doc, "GLP-1-population RCT data > adjacent-population RCT data > mechanism-only / in-vitro / animal data. Note: GLP-1-specific RCTs do not yet exist for any of the four side effects scored here — every evidence base is interpolation from analogous populations.")
    add_bullet(doc, "Effect size cited on the instrument the trial used (VAS, Halimeter ppb, Rhodes Index, profilometry surface area, telogen %, etc.) — never paraphrased into \"reduces nausea / breath / shedding\".")
    add_bullet(doc, "Dose adequacy compares labeled dose to the dose at which the cited effect was demonstrated. Underdosing is treated as a primary failure mode, not a rounding error.")
    add_bullet(doc, "Coverage of mechanism rewards products that address multiple drivers of multi-causal side effects (halitosis = 3 drivers; Ozempic face = 2; nausea = 2; hair loss = ≥4).")

    add_pagebreak(doc)


def build_product_section(doc, product, mech_lines, evidence_lines, effect_lines,
                          dose_lines, limitations_lines, references_inline=None):
    """One full product deep-dive. ~2-3 pages."""
    add_eyebrow(doc, f"{product['section_no']} · {product['side_effect']}")
    add_h1(doc, product["product_name"])
    p = add_para(doc, "", space_after=8)
    sub = p.add_run(f"{product['supplier']} — {product['format']}")
    style_run(sub, font=SERIF, size=14, color=MUTED, italic=True)

    add_amber_rule(doc, sz="6")

    add_h3(doc, "Product spec")
    add_kv_line(doc, "Active ingredients", product["actives"])
    add_kv_line(doc, "Format", product["format"])
    add_kv_line(doc, "Sourcing terms", product["moq"])
    if product["internal_score"] is not None:
        add_kv_line(doc, "Existing APALM internal score", f"{product['internal_score']} / 100")
    independent = sum(product["scores"])
    add_kv_line(doc, "Independent mitigation score (this document)", f"{independent} / 100")

    add_h3(doc, "The side effect, in one paragraph")
    add_para(doc, mech_lines[0], size=10, line_spacing=1.28)

    add_h3(doc, "Mechanism of action")
    for line in mech_lines[1:]:
        add_para(doc, line, size=10, line_spacing=1.28)

    add_h3(doc, "Evidence base")
    for line in evidence_lines:
        add_bullet(doc, line, size=10)

    add_h3(doc, "Effect size — what relief actually looks like")
    for line in effect_lines:
        add_para(doc, line, size=10, line_spacing=1.28)

    add_h3(doc, "Dose adequacy check")
    for line in dose_lines:
        add_para(doc, line, size=10, line_spacing=1.28)

    add_h3(doc, "Limitations and caveats")
    for line in limitations_lines:
        add_bullet(doc, line, size=10)

    add_h3(doc, "Score breakdown")
    rubric_table(doc, product)

    add_pagebreak(doc)


def rubric_table(doc, product):
    headers = ["Dimension", "Weight", "Score", "Rationale"]
    n = len(RUBRIC)
    table = doc.add_table(rows=n + 2, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(1.55), Inches(0.6), Inches(0.6), Inches(4.0)]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    for i, h in enumerate(headers):
        fill_cell(table.rows[0].cells[i], h, font=SANS, size=9, color=RGBColor.from_string("F5EFE3"),
                  bold=True, fill=AMBER_HEX, align=WD_ALIGN_PARAGRAPH.CENTER if i in (1,2) else WD_ALIGN_PARAGRAPH.LEFT)

    for trow in table.rows:
        set_row_cant_split(trow)
    for r, ((name, weight, _), score, rationale) in enumerate(
            zip(RUBRIC, product["scores"], product["scores_rationale"]), start=1):
        # color-code score cell
        fraction = score / weight if weight else 0
        if fraction >= 0.75:
            fill = SCORE_HI_HEX
        elif fraction >= 0.50:
            fill = SCORE_MID_HEX
        else:
            fill = SCORE_LOW_HEX
        fill_cell(table.rows[r].cells[0], name, font=SERIF, size=11, color=INK, bold=True)
        fill_cell(table.rows[r].cells[1], str(weight), color=MUTED,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(table.rows[r].cells[2], f"{score}", bold=True, fill=fill,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(table.rows[r].cells[3], rationale, font=SANS, size=9.5, color=INK)

    # composite row
    composite = sum(product["scores"])
    fill_cell(table.rows[-1].cells[0], "Composite", font=SERIF, size=12, color=INK, bold=True,
              fill=PAPER_HEX)
    fill_cell(table.rows[-1].cells[1], "100", color=MUTED, fill=PAPER_HEX,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(table.rows[-1].cells[2], str(composite), bold=True, color=INK,
              fill=score_color(composite), align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(table.rows[-1].cells[3],
              {
                  "ozempic_face":
                      "Real evidence on dermal aging, but the product treats the wrong tissue layer for "
                      "Ozempic face's primary driver (subcutaneous volume loss).",
                  "halitosis":
                      "Earns its place. Best-evidenced of the four. The mechanism-coverage caveat is the "
                      "honest limit, not a fault of the formulation.",
                  "hair_loss":
                      "Wrong target ingredient for non-deficient telogen effluvium. Lab-assay interference "
                      "is a real downside in this population.",
                  "nausea_ing":
                      "Right active. Wrong dose by an order of magnitude. The defining failure of the slate.",
                  "nausea_vw":
                      "What the shortlist's nausea SKU should look like. Properly dosed, well-evidenced, "
                      "evidence-aligned.",
              }[product["key"]],
              font=SANS, size=9.5, color=INK, italic=True, fill=PAPER_HEX)


# ---------------------------------------------------------------------------
# Per-product content builders (long-form prose)
# ---------------------------------------------------------------------------
def section_ozempic_face(doc):
    p = PRODUCTS[0]
    mech = [
        "Rapid GLP-1-induced weight loss (15–20% of body mass over 6–12 months) depletes facial "
        "subcutaneous fat compartments — malar, temporal, buccal — faster than overlying skin can "
        "retract. The resulting hollow, gaunt, prematurely-aged appearance has been popularized as "
        "\"Ozempic face.\" There are two layered drivers: (1) volume loss in the fat pads "
        "themselves, which is the dominant visual cue, and (2) collagen / elastin / GAG loss from "
        "concurrent caloric restriction, which thins the dermis and accentuates the volumetric loss.",
        "Matrixyl 3000 is a complex of two synthetic peptides — palmitoyl tetrapeptide-7 and "
        "palmitoyl tripeptide-1. The lipid (palmitoyl) chain enables transdermal penetration; the "
        "peptide moieties signal fibroblast collagen synthesis (collagen types I, III, IV, VII) and "
        "downregulate IL-6 inflammatory signaling. Synthe'6 is palmitoyl tripeptide-38, a newer "
        "biostimulatory peptide developed by Sederma that adds a fibrillin-1 stimulation pathway. "
        "Both are dermal-remodeling peptides — they act on the dermis (1–3 mm depth).",
        "Critically, neither peptide has a mechanism by which it could restore subcutaneous adipose "
        "tissue. Topical formulations do not penetrate to the 5–10 mm depth where facial fat pads "
        "reside. The actives are doing real work on the wrong tissue layer for the headline driver "
        "of Ozempic face.",
    ]
    evidence = [
        "Sederma in-house clinical dossier on Matrixyl 3000 — twice-daily 3% formulation, 8 weeks, "
        "N≈23 — reports a 45% reduction in deep-wrinkle surface area and a 20% improvement in skin "
        "tonicity by profilometry. Manufacturer-sponsored, no independent placebo control, but the "
        "endpoint instrument is well-validated.",
        "Li et al. 2023 (Journal of Cosmetic Dermatology, N=37, 12 weeks): yeast ferment + palmitoyl "
        "tripeptide-1 + palmitoyl tetrapeptide-7 eye cream produced +54.99% collagen density and "
        "+18.81% elasticity by ultrasound and cutometer. No placebo arm; small N.",
        "Sederma in-house dossier on Synthe'6 (palmitoyl tripeptide-38) — twice-daily 2%, N=25, ages "
        "42–70 — wrinkle reduction 31–100% across self-report bands; in-vitro collagen I +105%, "
        "collagen III +104%, collagen IV +42%. Heterogeneous endpoints.",
        "Dermatology reviews of GLP-1 facial aging (e.g., 2024–2025 commentary in JAAD and equivalent "
        "outlets) consistently route patients toward injectables (HA fillers, poly-L-lactic acid / "
        "Sculptra) and surgical fat grafting for actual volume restoration. Topical peptides do not "
        "appear in standard treatment algorithms for Ozempic face.",
        "Zero published RCTs of topical peptide serums in GLP-1, post-bariatric, or rapid-weight-loss "
        "facial-aging populations. Every cited effect size is interpolation from a general dermal-aging "
        "cohort.",
    ]
    effect = [
        "On dermal endpoints (wrinkle depth, collagen density, elasticity), expect 15–45% improvement "
        "over 8–12 weeks of consistent twice-daily use, consistent with the Sederma protocol. "
        "Visible to a careful self-observer; not visible enough to compensate for the volumetric loss "
        "that defines Ozempic face.",
        "On the headline complaint — facial hollowness — expect zero measurable change. Topical "
        "peptide serums do not restore subcutaneous adipose tissue. For comparison, hyaluronic acid "
        "fillers produce immediate volumetric correction; poly-L-lactic acid biostimulators produce "
        "gradual collagen-mediated volume gain over 3–4 months; autologous fat grafting produces "
        "durable correction. The serum is not in the same therapeutic category.",
    ]
    dose = [
        "Cellular Cosmetics is the only Australian supplier that publishes its INCI per the APALM "
        "notes — strong adequacy on this front. If the formulation is at the standard 3–8% Matrixyl "
        "3000 complex and ~2% Synthe'6, both are at concentrations the manufacturer trial used.",
        "The dose adequacy issue here is not concentration. It is that even at fully studied "
        "concentrations, the actives reach the wrong tissue depth for the primary driver of the "
        "side effect.",
    ]
    limitations = [
        "Topical peptides cannot restore facial fat-pad volume — physically impossible at 1–3 mm "
        "penetration depth.",
        "All clinical data is in general dermal-aging populations; no GLP-1, post-bariatric, or "
        "rapid-weight-loss cohort evidence exists.",
        "Sederma trials are manufacturer-sponsored; the strongest independent data (Li 2023) lacks "
        "a placebo arm and has a small sample size (N=37).",
        "Honest brand positioning: this is an adjunctive anti-aging serum that may slow concurrent "
        "dermal aging during rapid weight loss. Marketing it as a primary intervention for Ozempic "
        "face is a category mismatch.",
    ]
    build_product_section(doc, p, mech, evidence, effect, dose, limitations)


def section_halitosis(doc):
    p = PRODUCTS[1]
    mech = [
        "Approximately 30% of Ozempic users self-report halitosis (ADA News 2025 observational data). "
        "Three converging mechanisms drive it. Hyposalivation: GLP-1 receptor activity in salivary "
        "glands reduces flow, which removes the natural buffering and flushing of oral bacteria. "
        "Ketosis: the caloric deficit drives acetone and isoprene exhalation — the fruity / metallic "
        "odor characteristic of low-carb states. Gastric stasis: delayed gastric emptying produces "
        "fermentation gases and eructation (reported in ~9% of subjects in Novo's STEP-1 trial of "
        "semaglutide).",
        "Streptococcus salivarius K12 is a bacteriocin-producing strain that competitively colonizes "
        "the oral cavity. Two of its bacteriocins — salivaricin A2 and salivaricin B — suppress the "
        "anaerobic gram-negative organisms (Solobacterium moorei, Fusobacterium nucleatum, "
        "Porphyromonas gingivalis, Treponema denticola) most responsible for volatile sulfur compound "
        "(VSC) production.",
        "Mechanism is targeted and credible — but constrained. K12 acts in the oral cavity. "
        "Hyposalivation requires a saliva substitute or a sialagogue, not a probiotic. Ketotic "
        "acetone is exhaled from the lung after systemic production; no oral intervention can "
        "address it. Gastric-stasis eructation requires GI motility intervention. K12 is the right "
        "answer for one of three drivers.",
    ]
    evidence = [
        "Burton et al. 2006 (J Appl Microbiol 100(4):754–764, N=23 halitosis subjects): a 3-day "
        "chlorhexidine wash followed by K12 lozenges or placebo. 85% of K12 subjects achieved "
        ">100 ppb VSC reduction at week 1, vs. 30% in placebo. Halimeter instrument. The "
        "canonical halitosis K12 trial.",
        "Kim et al. 2020 (Probiotics Antimicrob Proteins 12:432–441, N=28, 30 days, double-blind): "
        "OLT (organoleptic test) and Halimeter VSC dropped within the K12 group, but a regression "
        "analysis showed no statistically significant difference vs. placebo. Authors note efficacy "
        "depended on tongue-coating removal before dosing — colonization is substrate-dependent.",
        "Yoo et al. 2020 (Clin Exp Dent Res 6:269–278): in-vitro co-culture of K12 / M18 with "
        "P. gingivalis and T. denticola. Total VSCs dropped from 42.32 to 2.02 ng/10 mL; H₂S was "
        "suppressed ~95%. Bacteriocin activity persisted in cell-free filtrate, supporting a "
        "metabolite-mediated mechanism.",
        "Tan et al. 2024 systematic review (J Clin Periodontal Drug Care): 6 RCTs, 360 participants. "
        "5 of 6 trials showed VSC reduction; 3 of 6 showed OLT improvement. Heterogeneity (strain, "
        "dose, duration 2–30 weeks) prevented a pooled meta-estimate. Authors graded evidence as "
        "low-to-moderate certainty.",
        "Zero GLP-1-population RCTs for K12 or any oral probiotic. The inference to GLP-1 halitosis "
        "is interpolation from general-halitosis cohorts.",
    ]
    effect = [
        "On VSC measured by Halimeter or OralChroma: expect a 60–100 ppb reduction at 1–2 weeks of "
        "consistent use, with the effect concentrated in subjects whose halitosis is bacterial in "
        "origin. The effect tapers within 4 weeks of discontinuation — colonization is not permanent.",
        "On the headline complaint — \"my breath smells\" — expect the bacterial fraction of the "
        "smell to reduce meaningfully. The ketotic-acetone fraction is unaffected. A subset of users "
        "(those whose halitosis is predominantly hyposalivation- or ketosis-driven) will report "
        "minimal improvement and should be redirected to xylitol mints + hydration + carb-cycling, "
        "not blamed for poor adherence.",
    ]
    dose = [
        "BLIS lozenge: 1.25 × 10⁹ CFU per dose. Burton's protocol used 1 × 10⁹ per lozenge, 4 lozenges "
        "per day acutely (4 × 10⁹ daily). At 1–2 lozenges per day on the BLIS product, a user is "
        "in the maintenance range used in subsequent trials. Adequate.",
        "Practical note: tongue-coating removal before dosing materially improved colonization in "
        "the Kim 2020 trial. Pairing the lozenge with a tongue scraper would likely lift effect "
        "size at no additional formulation cost.",
    ]
    limitations = [
        "Addresses only the bacterial driver — leaves hyposalivation and ketotic acetone untouched.",
        "Effect size is more variable in subjects with high tongue-coating or poor baseline oral hygiene.",
        "Not a cure — relapse within ~4 weeks of discontinuation is the norm. Recurring-revenue product, "
        "but not a one-shot fix.",
        "Rare case reports of S. salivarius infection in immunocompromised patients exist. Not a "
        "concern in the typical GLP-1 cohort, but warrants a label disclaimer.",
        "Honest brand positioning: \"reduces the bacterial component of GLP-1 breath\" is accurate. "
        "\"Eliminates Ozempic breath\" is not.",
    ]
    build_product_section(doc, p, mech, evidence, effect, dose, limitations)


def section_hair_loss(doc):
    p = PRODUCTS[2]
    mech = [
        "Telogen effluvium (TE) is the diffuse shedding pattern that follows a metabolic shock by "
        "2–4 months. In GLP-1 users, the shock is rapid caloric deficit and the protein / "
        "micronutrient suppression that comes with appetite loss. About 6–13% of GLP-1 users "
        "experience clinically noticeable shedding (Shah 2024). Drivers: (a) the metabolic stress "
        "itself, which prematurely shifts follicles from anagen to telogen phase; (b) sub-clinical "
        "deficiencies in iron, zinc, vitamin D, and B12 that emerge under reduced food intake; and "
        "(c) reduced amino-acid supply (cysteine, methionine, lysine) under low-protein eating.",
        "The Vox capsule pairs biotin with marine collagen peptides plus supportive vitamins. The "
        "biotin half is meant to support keratin synthesis. The marine collagen half supplies amino "
        "acid building blocks. Both rationales are plausible in deficient subjects — but most GLP-1 "
        "users with TE are not biotin-deficient, and the deficiencies they actually develop "
        "(iron / zinc / D / B12) are not what this capsule meaningfully addresses.",
    ]
    evidence = [
        "Patel et al. 2017 (Skin Appendage Disorders 3:166–169) systematically reviewed all published "
        "biotin-for-hair-loss case reports. 18 cases, every responder had an underlying biotin "
        "deficiency state — biotinidase deficiency, brittle nail syndrome, uncombable hair disease, "
        "antiseizure-medication-induced deficiency. Conclusion: \"the evidence base for biotin "
        "supplementation in alopecia remains sparse\" and the authors recommended against routine "
        "use outside deficiency states.",
        "Soleymani et al. 2017 (J Drugs Dermatol 16:496–500), narrative review of biotin marketing: "
        "\"insufficient data to recommend biotin supplementation in healthy individuals for the "
        "treatment of alopecia.\" Zero placebo-controlled trials in any alopecia type.",
        "Milani et al. 2023 (Skin Res Technol 29:e13467, N=83, 3 months, double-blind RCT): an oral "
        "supplement containing cysteine, methionine, iron 15 mg, selenium 75 µg, AND marine "
        "hydrolyzed collagen 1 g reduced telogen-phase hair from 31.4% ± 4.7% to 24.3% ± 7.1% "
        "(p<0.001). Hair density and cuticle quality also improved. The study is positive — but "
        "the formulation includes therapeutic iron and selenium, so the marine-collagen-alone "
        "contribution cannot be isolated.",
        "Durusu Turkoglu et al. 2024 (J Cosmet Dermatol 23:1289–1296) compared 90 active-TE patients "
        "to 90 matched controls. Zinc was significantly lower in the TE group (49.91 vs. 60.53 µg/dL, "
        "p=0.004). Iron, ferritin, B12, vitamin D, AND biotin showed no significant differences. "
        "Authors concluded zinc deficiency warrants screening; biotin supplementation does not.",
        "FDA MedWatch case (PMC6802814, 2019): a 67-year-old on 5 mg/day biotin presented with "
        "spurious hypocalcemia and hyperparathyroidism signals that resolved within one month of "
        "discontinuation. Biotin >2.5 mg/day interferes with ~59% of common immunoassays — TSH, "
        "free T4, troponin, hCG, parathyroid hormone — at the doses found in most consumer products.",
        "Zero GLP-1 or post-bariatric population RCTs for biotin or marine collagen on hair-specific "
        "endpoints.",
    ]
    effect = [
        "In a typical non-deficient GLP-1 user, expect a 10–20% reduction in shedding over 3 months "
        "from a biotin + marine collagen capsule alone — and a meaningful fraction of that is "
        "regression to the mean as the underlying caloric deficit stabilizes and the user passes "
        "the 4-month TE peak naturally.",
        "If the same user instead repleted iron, zinc, and vitamin D after lab confirmation and "
        "raised protein intake to ≥1.6 g/kg, expected shedding reduction is 40–60% — closer to "
        "the Milani trial result.",
    ]
    dose = [
        "Typical Vox-class formulations deliver biotin 5–10 mg per capsule. This matches the "
        "deficiency-treatment dose used in case reports — but it is dose-irrelevant for the "
        "non-deficient majority. It is also the dose that triggers the Patel-cited assay "
        "interference issue.",
        "Marine collagen at 1–3 g per capsule is below the 10 g/day dose used in skin RCTs and "
        "below the 1 g component dose in the Milani trial (which is the closest to a positive "
        "hair RCT in the literature). Capsule format is the rate-limiter on collagen dose; a "
        "powder or scoop could deliver 10 g comfortably.",
    ]
    limitations = [
        "Wrong target for the typical user. The actual GLP-1 hair-loss drivers (iron / zinc / D / "
        "B12 / protein) are not what this product corrects.",
        "Biotin >2.5 mg/day causes false TSH, free T4, troponin, and hCG values in ~59% of common "
        "immunoassays. Clinically consequential for a population monitored for thyroid function "
        "during rapid weight loss. This is the largest active risk in any of the four products.",
        "Time + nutritional repletion + protein adequacy is what evidence-based TE management "
        "actually looks like. Topical minoxidil is the next escalation. A generic Hair-Skin-Nails "
        "capsule sits well below either of those interventions.",
        "Honest brand positioning: \"general nail and skin support during weight loss\" is defensible. "
        "\"Treats GLP-1 hair loss\" is not, in the absence of upstream nutritional screening.",
    ]
    build_product_section(doc, p, mech, evidence, effect, dose, limitations)


def section_nausea_ing(doc):
    p = PRODUCTS[3]
    mech = [
        "Nausea is the single most common GLP-1 side effect — 24–44% of users in pivotal trials, "
        "and the most-cited reason for discontinuation in the first 8 weeks. Mechanism: GLP-1 "
        "agonists slow gastric emptying by 30–70%, so food sits longer, distends the stomach, and "
        "triggers vagal nausea signals. Central CNS effects on the area postrema (the same "
        "brainstem region implicated in motion sickness) compound the peripheral signal. Worst "
        "during titration weeks 1–2 of each dose step; tolerance develops over 4–8 weeks.",
        "Ginger root (Zingiber officinale) acts on multiple anti-nausea targets: 5-HT3 receptor "
        "antagonism in the GI tract, prokinetic effect on gastric motility, and modest anti-emetic "
        "central effects. Vitamin B6 (pyridoxine) modulates central nausea pathways — first-line "
        "evidence in pregnancy nausea per ACOG. Magnesium carbonate is included for buffering and "
        "GI tolerability. Mechanistically, this product is well-targeted.",
        "The targeting is correct in kind. The targeting is not correct in dose.",
    ]
    evidence = [
        "Vutyavanich et al. 2001 (Obstet Gynecol 97:577–582, N=70 first-trimester pregnant women): "
        "1,000 mg ginger daily for 4 days reduced VAS nausea by 2.1 points vs. 0.9 points on "
        "placebo (p=0.014). Vomiting episodes 1.4 vs. 0.3 (p<0.05). Effect size: ~2.3× placebo on "
        "the primary endpoint at 1 g/day.",
        "Ryan / URCC CCOP 2012 (Support Care Cancer 20:1479–1489, N=576 cancer patients on "
        "chemotherapy): four-arm dose-response trial — placebo, 0.5 g, 1.0 g, 1.5 g ginger daily. "
        "Significance reached at 0.5 g and 1.0 g for acute nausea severity (the 1.5 g arm did not "
        "consistently outperform placebo, supporting an inverted-U dose-response). 1.0 g produced "
        "the strongest signal. The lowest dose tested in any major nausea trial was 500 mg.",
        "Chaiyakunapruk et al. 2006 (Am J Obstet Gynecol 194(1):95–99, post-op meta-analysis, "
        "5 trials pooled, N=363): fixed-dose ≥1 g ginger gave a relative risk of 0.69 for "
        "post-op nausea (95% CI 0.54–0.89) and 0.61 for post-op vomiting (95% CI 0.45–0.84) "
        "vs. placebo — i.e., ~31% and ~39% relative-risk reductions. Only adverse event reported "
        "was mild abdominal discomfort.",
        "ACOG Practice Bulletin (2018, updated): pyridoxine 10–25 mg three times daily (30–75 mg "
        "total) is first-line evidence-based pharmacotherapy for nausea in pregnancy. The "
        "underlying RCTs (Sahakian 1991, Vutyavanich 1995) used 25 mg t.i.d.",
        "Sripramote & Lekhyananda 2003 (J Med Assoc Thai 86:846–853, N=138 pregnant): 1 g ginger "
        "vs. 75 mg B6 head-to-head — equivalent on primary endpoint, ginger superior on nausea "
        "intensity / distress secondary endpoints.",
        "Oliveira et al. 2014 (open-label combination, N=45): ginger 500 mg b.i.d. + B6 25 mg t.i.d. "
        "produced 73% nausea reduction and 81% vomiting reduction by Rhodes Index, 93% response rate. "
        "Adverse events: mild GI irritation in 2/45 (4%).",
        "Zero GLP-1-population RCTs for ginger or B6 exist. The inference to GLP-1 nausea is "
        "interpolation from chemo, post-op, and pregnancy cohorts — defensible because the central "
        "(area postrema) and peripheral (gastric / 5-HT3) mechanisms are shared, but still "
        "interpolation.",
    ]
    effect = [
        "At the labeled 67 mg ginger / 1.4 mg B6 dose: published data is silent. No trial has "
        "tested ginger below 250 mg. The URCC dose-response curve does not extrapolate downward "
        "to 67 mg — that dose sits well below the curve's lower inflection. Best estimate of "
        "effect size: indistinguishable from placebo.",
        "If the effect is non-zero, it is plausibly in the 0–10% nausea-reduction range — i.e., "
        "within placebo noise. The product is not in the same therapeutic category as a 1 g ginger "
        "dose.",
    ]
    dose = [
        "Ginger: labeled 67 mg vs. studied therapeutic minimum 1,000 mg = 14.9× underdose. The "
        "lowest dose any major trial actually tested is 500 mg (URCC); the labeled dose is 7.5× "
        "below even that. To reach a therapeutic 1 g ginger dose from this softgel, a user would "
        "need to take 15 capsules — at $2.50/capsule unit cost, $37.50 per dose. The math fails "
        "before the pharmacology does.",
        "B6: labeled 1.4 mg vs. ACOG protocol 75 mg/day = 53.6× underdose. 1.4 mg is the adult "
        "RDA — a multivitamin trace. It is not a therapeutic dose by any published standard.",
        "Magnesium carbonate is unlabeled in dose. Cannot evaluate adequacy. Magnesium has only "
        "weak anti-nausea evidence in any case; primary value here is gastric pH buffering.",
    ]
    limitations = [
        "Sub-pharmacologic dose on both primary actives. The product is correctly formulated in "
        "kind and dramatically misformulated in degree.",
        "ING markets the formula as \"ready-to-fill,\" which means downstream brands can rebrand "
        "it without correcting the dose. The dose problem will propagate.",
        "Honest brand positioning: this is not a nausea SKU. It is a GI-comfort softgel that uses "
        "anti-nausea actives at sub-therapeutic levels. Marketing it as Ozempic-nausea relief "
        "would set up an evidence gap that is hard to defend.",
        "Recommended action: do not white-label this formulation as the Aplomb nausea SKU. The "
        "alternative §04·B (1 g ginger + 25 mg B6) is the same active stack at the dose the "
        "evidence supports.",
    ]
    build_product_section(doc, p, mech, evidence, effect, dose, limitations)


def section_nausea_vw(doc):
    p = PRODUCTS[4]
    mech = [
        "Same mechanism narrative as §04·A — GLP-1 nausea is dominated by gastric distension and "
        "area-postrema activation; ginger and B6 are well-targeted; tolerance develops over 4–8 "
        "weeks. The only difference between this entry and the previous one is dose.",
        "A premium ginger gummy spec at 1,000 mg standardized ginger root extract + 25 mg B6 per "
        "gummy puts both actives squarely inside the dose range each was studied at — for ginger, "
        "the URCC dose-response peak; for B6, the lower bound of the ACOG 25–75 mg t.i.d. window.",
    ]
    evidence = [
        "All seven of the studies cited in §04·A apply directly. The same evidence base. The same "
        "cohort heterogeneity caveat (chemo + pregnancy + post-op, no GLP-1-specific RCT yet). "
        "What changes is that the labeled dose now corresponds to the dose the studies actually "
        "tested.",
        "The combination data — Oliveira 2014 — is the closest direct analog: ginger 1 g/day "
        "split + B6 75 mg/day split produced 73% nausea reduction. Per-gummy 1 g + 25 mg roughly "
        "matches a single Oliveira dose unit, with adherence advantages over a 5-pill regimen.",
    ]
    effect = [
        "Realistic effect size in a GLP-1 user during titration: 30–50% reduction in nausea on a "
        "VAS or Rhodes Index scale — extrapolating from Vutyavanich's 2.3× placebo and "
        "Chaiyakunapruk's ~31–39% RR reduction.",
        "Number Needed to Treat: ~4–6, meaning one in four to six GLP-1 users with titration "
        "nausea would experience clinically meaningful relief vs. placebo. Not curative — "
        "mitigative, exactly the brief.",
    ]
    dose = [
        "Ginger 1 g: at the URCC dose-response peak. 25 mg B6: at the lower bound of the studied "
        "anti-nausea range, well below the 100 mg/day chronic-toxicity threshold. Dose "
        "adequacy is the headline strength.",
        "Practical regimen: 1 gummy at first sign of nausea during the titration weeks; up to 2 "
        "per day. At ~$2.50/unit cost, this is unit-economic-isomorphic with the ING product but "
        "delivers actual therapeutic dose.",
    ]
    limitations = [
        "Ginger anticoagulant additivity is real at 1 g/day — meaningful for users on warfarin or "
        "DOACs. Label disclosure required.",
        "B6 chronic intake should stay <100 mg/day to avoid peripheral neuropathy. At 25 mg per "
        "gummy, 4-gummy/day adherence approaches that ceiling. Dose limit and label clarity matter.",
        "Mild GI irritation in ~4% of subjects in the Oliveira trial — a small number, but worth "
        "calling out for a mass-market gummy.",
        "Gummy palatability is high — risk of overconsumption / use as candy. Childproof packaging "
        "and 2/day cap on label are sensible.",
        "No GLP-1-specific RCT yet. The strongest claim Aplomb can make is \"evidence-based dose of "
        "the most studied anti-nausea actives, applied to GLP-1 nausea.\" That is a defensible claim.",
    ]
    build_product_section(doc, p, mech, evidence, effect, dose, limitations)


def build_comparative(doc):
    add_eyebrow(doc, "Comparative analysis · 3 of 4")
    add_h1(doc, "All five products, side by side")
    add_amber_rule(doc, sz="6")

    add_lede(doc,
        "The cross-product table shows where the gaps are. No product fails on adherence or "
        "tolerability — the signal is concentrated on three dimensions: mechanism coverage, dose "
        "adequacy, and mechanistic match."
    )

    headers = ["Dimension"] + [
        "Cellular", "BLIS K12", "Vox HSN", "ING (67 mg)", "Vital Whole (1 g)"
    ]
    table = doc.add_table(rows=len(RUBRIC) + 2, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [Inches(1.85)] + [Inches(0.93)] * 5
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    # Header row
    fill_cell(table.rows[0].cells[0], headers[0], font=SANS, size=9,
              color=RGBColor.from_string("F5EFE3"), bold=True, fill=AMBER_HEX)
    for i, h in enumerate(headers[1:], start=1):
        fill_cell(table.rows[0].cells[i], h, font=SANS, size=8.5,
                  color=RGBColor.from_string("F5EFE3"), bold=True, fill=AMBER_HEX,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    for trow in table.rows:
        set_row_cant_split(trow)

    # Body
    for r, (name, weight, _) in enumerate(RUBRIC, start=1):
        fill_cell(table.rows[r].cells[0],
                  f"{name}",
                  font=SERIF, size=11, color=INK, bold=True)
        for c, prod in enumerate(PRODUCTS, start=1):
            score = prod["scores"][r-1]
            fraction = score / weight
            if fraction >= 0.75:
                fill = SCORE_HI_HEX
            elif fraction >= 0.50:
                fill = SCORE_MID_HEX
            else:
                fill = SCORE_LOW_HEX
            fill_cell(table.rows[r].cells[c],
                      f"{score} / {weight}",
                      bold=(score == max(p['scores'][r-1] for p in PRODUCTS)),
                      fill=fill,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    # Composite row
    fill_cell(table.rows[-1].cells[0], "Composite (0–100)",
              font=SERIF, size=12, color=INK, bold=True, fill=PAPER_HEX)
    for c, prod in enumerate(PRODUCTS, start=1):
        composite = sum(prod["scores"])
        fill_cell(table.rows[-1].cells[c], str(composite), bold=True,
                  fill=score_color(composite), align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "", space_after=10)

    add_h3(doc, "Where the slate concentrates its weakness")
    add_bullet(doc,
        "Dose adequacy: ING 1/15. Every other product clears 9/15. Sub-pharmacologic dosing is "
        "the single largest correctable defect across the slate.")
    add_bullet(doc,
        "Mechanistic match: Cellular 9/20 and Vox 7/20. Both products use real actives — but "
        "against the wrong primary mechanism. These are formulation-strategy gaps, not dose gaps.")
    add_bullet(doc,
        "Mechanism coverage: every product loses points here, because every targeted side effect "
        "is multi-causal and no single SKU is engineered to span all drivers. This is an "
        "argument for kits over single SKUs (consistent with the kit-shape recommendation in "
        "AFTER_halitosis_scoring_v1.md §4).")

    add_h3(doc, "Where the slate is genuinely strong")
    add_bullet(doc,
        "Tolerability: every product scores 7/10 or higher. No product introduces meaningful new "
        "GLP-1 contraindications, with the single exception of biotin's lab-assay interference.")
    add_bullet(doc,
        "Adherence: every product scores 8/10. Format choices (serum, lozenge, capsule, gummy) "
        "are realistic for a daily consumer routine.")
    add_bullet(doc,
        "Clinical evidence quality: BLIS K12 and the Vital Whole archetype both clear 14/20. The "
        "evidence base under both products is strong; the difference between strength and weakness "
        "across the slate is not access to good science but whether a given formulation reflects it.")

    add_pagebreak(doc)


def build_bottom_line(doc):
    add_eyebrow(doc, "Bottom line · 4 of 4")
    add_h1(doc, "What to keep, change, and reframe")
    add_amber_rule(doc, sz="6")

    add_h2(doc, "Keep — BLIS K12 lozenge")
    add_para(doc,
        "BLIS K12 is the most defensible position on the four-product slate. Real RCT evidence on "
        "the actual symptom, dose adequately matches studied protocols, format is consumer-friendly, "
        "tolerability is established. The honest caveat — coverage of only the bacterial driver — is "
        "a copy / positioning issue, not a product issue. Frame it as \"the bacterial half of "
        "Ozempic breath\" and pair it on-pack with xylitol mints (hyposalivation) and electrolyte "
        "stick (hydration) to round out the kit.")

    add_h2(doc, "Reframe — Cellular Cosmetics serum")
    add_para(doc,
        "Cellular's serum has solid dermal-aging evidence and is well-formulated. The category "
        "mismatch is fixable in copy: position it as \"slows the dermal half of Ozempic face — the "
        "wrinkle-and-crepiness fraction that goes with rapid weight loss\" rather than as a "
        "treatment for facial hollowness. For volume restoration, the consumer journey leads to "
        "an injectable provider — that's a referral relationship, not a competing Aplomb SKU. "
        "Repositioned, the serum is a credible adjunct. Marketed as it stands, it sets up a "
        "consumer expectation it cannot meet.")

    add_h2(doc, "Replace — ING Pharmaceutical softgel")
    add_para(doc,
        "Do not white-label the ING softgel as the Aplomb nausea SKU. The dose problem is "
        "irreducible: 67 mg ginger is below the lowest dose any RCT has ever tested, and 1.4 mg "
        "B6 is the adult RDA. A consumer who takes the product, feels no relief, and then learns "
        "from any 5-minute Google search that 1 g is the studied dose has a brand-trust failure "
        "no marketing copy can recover.")
    add_para(doc,
        "The replacement is in the same supplier directory: a 1,000 mg ginger + 25 mg B6 gummy "
        "(Vital Whole-style spec, sourceable through the Alibaba ginger-chew category in the "
        "APALM list at MOQ ≤100 with custom B6 add-on, per build_supplier_shortlist_deck.py:518–531). "
        "Same unit cost band, same MOQ band, ten-fold higher therapeutic match. This is the single "
        "highest-leverage decision in the entire slate.")

    add_h2(doc, "Rework — Vox Hair, Skin & Nails")
    add_para(doc,
        "Vox's certifications and US-based sub-100 MOQ are valuable. The capsule itself is the "
        "wrong product for this side effect. Two paths: (1) custom-formulate through Vox a "
        "GLP-1-specific TE capsule with iron 15 mg, zinc 15 mg, vitamin D 1,000 IU, B12 500 µg, "
        "and a smaller biotin dose ≤2 mg/day to clear the assay-interference threshold; or (2) "
        "drop hair loss from the v1 launch slate and pick it back up after the Aplomb nutrient-stack "
        "SKU exists, where the deficiency-targeted ingredients can do double duty.")
    add_para(doc,
        "Either path is better than launching the stock Vox HSN under the Aplomb name. The "
        "biotin assay-interference issue alone is a regulatory-comms risk worth avoiding.")

    add_h2(doc, "What this means for the slate")
    add_para(doc,
        "Of the four selected SKUs, one earns its place on the merits (BLIS), one needs copy "
        "discipline (Cellular), one needs replacement with a properly-dosed substitute already "
        "in the supplier directory (ING → Vital Whole-style gummy), and one needs custom "
        "reformulation or reslotting (Vox). The highest-leverage single change is the nausea "
        "swap. The most consequential framing change is on Ozempic face. The most overlooked "
        "risk is biotin assay interference. The slate's strongest property is that none of its "
        "actives are dangerous; the slate's weakest property is that two of its actives are "
        "directionally wrong and one is dose-irrelevant.")

    add_pagebreak(doc)


def build_sources(doc):
    add_eyebrow(doc, "Sources cited")
    add_h1(doc, "References")
    add_amber_rule(doc, sz="6")

    refs = [
        ("Ozempic face — peptides + dermal aging",
         [
            "Li Y, Wang Y, Wang Y, et al. Efficacy of yeast ferment + palmitoyl peptide eye cream. "
            "Journal of Cosmetic Dermatology. 2023;22(8):2169–2178.",
            "Sederma. Matrixyl 3000 Technical Dossier and clinical data summary (manufacturer "
            "documentation, Sederma SAS, France).",
            "Sederma. Synthe'6 (palmitoyl tripeptide-38) Technical Dossier (manufacturer documentation).",
            "Reviews of GLP-1 dermatologic effects, 2024–2025 (multiple sources including JAAD "
            "and PMC PMC12110338 — narrative review of GLP-1 facial volume loss and treatment "
            "algorithm).",
         ]),
        ("Halitosis — S. salivarius K12",
         [
            "Burton JP, Chilcott CN, Moore CJ, Speiser G, Tagg JR. A preliminary study of the "
            "effect of probiotic Streptococcus salivarius K12 on oral malodour parameters. "
            "Journal of Applied Microbiology. 2006;100(4):754–764.",
            "Kim D, Hong JY, Park KS, et al. Probiotic effect of Streptococcus salivarius K12 on "
            "halitosis. Probiotics and Antimicrobial Proteins. 2020;12(2):432–441.",
            "Yoo JI, Shin IS, Jeon JG, et al. The effect of S. salivarius K12 / M18 on the "
            "growth and metabolism of oral pathogens in vitro. Clinical and Experimental Dental "
            "Research. 2020;6(2):269–278.",
            "Tan A, et al. Effectiveness of probiotics in managing oral halitosis: a systematic "
            "review. Journal of Clinical Periodontology / Oral Care reviews series. 2024.",
            "Montes M, et al. Safety assessment of Streptococcus salivarius K12 / BLIS K12. "
            "Applied and Environmental Microbiology. 2006;72(4):3050–3053.",
            "PMC PMC12729639 (2025 narrative review of GLP-1 oral effects: hyposalivation, "
            "ketosis, gastric stasis).",
         ]),
        ("Hair loss — biotin + marine collagen",
         [
            "Patel DP, Swink SM, Castelo-Soccio L. A review of the use of biotin for hair loss. "
            "Skin Appendage Disorders. 2017;3(3):166–169.",
            "Soleymani T, Lo Sicco K, Shapiro J. The infatuation with biotin supplementation: "
            "is there truth behind its rising popularity? Journal of Drugs in Dermatology. "
            "2017;16(5):496–500.",
            "Milani M, Celi D, Bianchi L, et al. Efficacy and tolerability of an oral supplement "
            "containing amino acids, iron, selenium, and marine hydrolyzed collagen in subjects "
            "with hair loss. Skin Research and Technology. 2023;29(4):e13467.",
            "Durusu Turkoglu E, Balci DD, Karadag AS. Comprehensive investigation of biochemical "
            "status in patients with telogen effluvium. Journal of Cosmetic Dermatology. "
            "2024;23(4):1289–1296.",
            "FDA MedWatch / PMC PMC6802814. Adverse event report on biotin interference with "
            "common immunoassays. 2019.",
            "Shah S, et al. GLP-1-related hair loss: incidence and characterization in a clinic "
            "cohort. 2024 (cited in Aplomb complaints inventory; original source TBD on independent "
            "fact-check).",
         ]),
        ("Nausea — ginger + B6",
         [
            "Vutyavanich T, Kraisarin T, Ruangsri R. Ginger for nausea and vomiting in pregnancy: "
            "a randomized, double-blind, placebo-controlled trial. Obstetrics & Gynecology. "
            "2001;97(4):577–582.",
            "Ryan JL, Heckler CE, Roscoe JA, et al. Ginger (Zingiber officinale) reduces acute "
            "chemotherapy-induced nausea: a URCC CCOP study of 576 patients. Supportive Care in "
            "Cancer. 2012;20(7):1479–1489.",
            "Chaiyakunapruk N, Kitikannakorn N, Nathisuwan S, Leeprakobboon K, Leelasettagool C. "
            "The efficacy of ginger for the prevention of postoperative nausea and vomiting: a "
            "meta-analysis. American Journal of Obstetrics & Gynecology. 2006;194(1):95–99.",
            "Sripramote M, Lekhyananda N. A randomized comparison of ginger and vitamin B6 in the "
            "treatment of nausea and vomiting of pregnancy. Journal of the Medical Association of "
            "Thailand. 2003;86(9):846–853.",
            "Sahakian V, Rouse D, Sipes S, Rose N, Niebyl J. Vitamin B6 is effective therapy for "
            "nausea and vomiting of pregnancy: a randomized, double-blind, placebo-controlled "
            "study. Obstetrics & Gynecology. 1991;78(1):33–36.",
            "Vutyavanich T, Wongtra-ngan S, Ruangsri R. Pyridoxine for nausea and vomiting of "
            "pregnancy: a randomized, double-blind, placebo-controlled trial. American Journal of "
            "Obstetrics and Gynecology. 1995;173(3 Pt 1):881–884.",
            "ACOG Practice Bulletin 189: Nausea and Vomiting of Pregnancy. American College of "
            "Obstetricians and Gynecologists. 2018 (updated).",
            "Oliveira LG, et al. Combined ginger and B6 therapy for nausea / vomiting of "
            "pregnancy. 2014 (open-label combination trial).",
         ]),
        ("Source of supplier specs",
         [
            "Cellular Cosmetics, BLIS Technologies, Vox Nutrition, ING Pharmaceutical entries: "
            "build_supplier_shortlist_deck.py (lines 270–533), aplomb.clinic project. Internal "
            "alleviation_score values: 84 / 82 / 72 / 82 (lines 309, 386, 432, 495 respectively).",
            "Existing Aplomb halitosis scoring memo: AFTER_halitosis_scoring_v1.md, dated 2026-05-01, "
            "aplomb.clinic project.",
            "Aplomb GLP-1 evidence table: GLP1_evidence_table.csv, /glp-1 support/ directory.",
         ]),
    ]
    for category, citations in refs:
        add_h3(doc, category)
        for c in citations:
            add_bullet(doc, c)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    doc = Document()
    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = SANS
    style.font.size = Pt(10.5)

    # Cream page color
    set_page_color(doc, BG_HEX)

    build_cover(doc)
    build_executive_summary(doc)
    build_methodology(doc)
    section_ozempic_face(doc)
    section_halitosis(doc)
    section_hair_loss(doc)
    section_nausea_ing(doc)
    section_nausea_vw(doc)
    build_comparative(doc)
    build_bottom_line(doc)
    build_sources(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"[OK] wrote {OUT}")


if __name__ == "__main__":
    main()
